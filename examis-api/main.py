from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, HttpUrl
from typing import List, Optional
import requests
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from PIL import Image, ImageOps

app = FastAPI(title="Examis AI Document Generator")

# --- 1. PYDANTIC MODELS (Payload Validation) ---
class MarksData(BaseModel):
    mcq_points: int
    short_points: int
    long_points: int
    fib_points: Optional[int] = 1

class CustomScenarioItem(BaseModel):
    text: str
    marks: Optional[int] = 0

class MCQItem(BaseModel):
    question: str
    options: List[str]
    target_clo: Optional[str] = None
    image_url: Optional[HttpUrl] = None

class ShortQuestionItem(BaseModel):
    question: str
    target_clo: Optional[str] = None
    image_url: Optional[HttpUrl] = None

class LongQuestionItem(BaseModel):
    question: str
    target_clo: Optional[str] = None
    image_url: Optional[HttpUrl] = None

class FillInTheBlankItem(BaseModel):
    question: str
    answer: str 
    target_clo: Optional[str] = None
    image_url: Optional[HttpUrl] = None

class DiagramQuestionItem(BaseModel):
    question: str
    image_url: HttpUrl
    target_clo: Optional[str] = None
    marks: Optional[int] = 0

class ExamData(BaseModel):
    title: str
    marks: MarksData
    custom_scenarios: List[CustomScenarioItem] = []
    mcqs: List[MCQItem] = []
    fillInTheBlanks: List[FillInTheBlankItem] = []
    shortQuestions: List[ShortQuestionItem] = []
    longQuestions: List[LongQuestionItem] = []
    diagram_questions: List[DiagramQuestionItem] = [] 


class DocumentRequest(BaseModel):
    template_url: HttpUrl
    show_clo_tags: bool = False
    exam_data: ExamData


# --- 2. CORE LOGIC ---
def process_exam(payload: DocumentRequest) -> io.BytesIO:
    response = requests.get(str(payload.template_url))
    if response.status_code != 200:
        raise HTTPException(status_code=400, detail="Failed to download template")
    
    doc_stream = io.BytesIO(response.content)
    doc = Document(doc_stream)
    exam = payload.exam_data
    show_clo = payload.show_clo_tags

    # Locate the anchor point and add the italic instructions
    for p in doc.paragraphs:
        if "{{START_EXAM_HERE}}" in p.text:
            p.text = p.text.replace("{{START_EXAM_HERE}}", "")
            instructions = "[Encircle the correct options. Overwriting will not be entertained. Multiple answers in fill in the blanks will be considered void.]"
            inst_run = p.add_run(instructions)
            inst_run.italic = True

    def add_section_header(title: str, points: int, count: int):
        p = doc.add_paragraph()
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        marks_str = f"[{points} x {count}]"
        p.add_run(title).bold = True
        p.add_run(f"\t{marks_str}").bold = True

    def insert_image_if_exists(img_url):
        if img_url:
            try:
                img_response = requests.get(str(img_url))
                if img_response.status_code == 200:
                    img_stream = io.BytesIO(img_response.content)
                    
                    # --- 1. THE EXIF ROTATION FIX ---
                    # Open the image with Pillow
                    img = Image.open(img_stream)
                    
                    # Automatically rotate it upright based on phone camera EXIF data
                    img = ImageOps.exif_transpose(img)
                    
                    # Prevent color mode errors
                    if img.mode in ("RGBA", "P"):
                        img = img.convert("RGB")
                        
                    # Save the fixed, upright image back to a new RAM stream
                    fixed_stream = io.BytesIO()
                    img.save(fixed_stream, format='PNG')
                    fixed_stream.seek(0)

                    # --- 2. THE ALIGNMENT FIX ---
                    # Create a brand new, clean paragraph just for the image
                    img_paragraph = doc.add_paragraph()
                    
                    # Force the paragraph to align strictly to the LEFT 
                    img_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Add the fixed picture
                    img_run = img_paragraph.add_run()
                    img_run.add_picture(fixed_stream, width=Inches(4.5))
                    
            except Exception as e:
                print(f"Warning: Failed to load image - {e}")

    # Write Custom Scenarios
    if exam.custom_scenarios:
        scenarios_header = doc.add_paragraph()
        scenarios_header.add_run("Case Studies / Scenarios").bold = True
        for i, scenario in enumerate(exam.custom_scenarios, 1):
            p = doc.add_paragraph()
            p.add_run(f"Scenario {i} ({scenario.marks} Marks)").bold = True
            doc.add_paragraph(scenario.text)
        doc.add_paragraph()

    # Write MCQs
    if exam.mcqs:
        add_section_header("Multiple Choice Questions", exam.marks.mcq_points, len(exam.mcqs))
        for i, mcq in enumerate(exam.mcqs, 1):
            p = doc.add_paragraph()
            q_text = f"{i}. {mcq.question}"
            if show_clo and mcq.target_clo:
                q_text += f" [{mcq.target_clo}]"
            p.add_run(q_text).bold = True
            insert_image_if_exists(mcq.image_url)
            
            # Formats options on the same line using spaces
            labels = ['a)', 'b)', 'c)', 'd)']
            formatted_options = "    ".join([f"{labels[idx]} {opt}" for idx, opt in enumerate(mcq.options) if idx < len(labels)])
            opt_p = doc.add_paragraph(formatted_options)
            opt_p.paragraph_format.left_indent = Inches(0.5)
        doc.add_paragraph() 

    # Write Fill in the Blanks
    if exam.fillInTheBlanks:
        add_section_header("Fill in the Blanks", exam.marks.fib_points, len(exam.fillInTheBlanks))
        for i, fib in enumerate(exam.fillInTheBlanks, 1):
            p = doc.add_paragraph()
            q_text = f"{i}. {fib.question}"
            if show_clo and fib.target_clo:
                q_text += f" [{fib.target_clo}]"
            p.add_run(q_text).bold = True
            insert_image_if_exists(fib.image_url)
        doc.add_paragraph()

    # Write Short Questions
    if exam.shortQuestions:
        add_section_header("Short Answer Questions", exam.marks.short_points, len(exam.shortQuestions))
        for i, sq in enumerate(exam.shortQuestions, 1):
            p = doc.add_paragraph()
            q_text = f"{i}. {sq.question}"
            if show_clo and sq.target_clo:
                q_text += f" [{sq.target_clo}]"
            p.add_run(q_text).bold = True
            insert_image_if_exists(sq.image_url)
            doc.add_paragraph()  

    # Write Long Questions
    if exam.longQuestions:
        add_section_header("Long Answer Questions", exam.marks.long_points, len(exam.longQuestions))
        for i, lq in enumerate(exam.longQuestions, 1):
            p = doc.add_paragraph()
            q_text = f"{i}. {lq.question}"
            if show_clo and lq.target_clo:
                q_text += f" [{lq.target_clo}]"
            p.add_run(q_text).bold = True
            insert_image_if_exists(lq.image_url)
            doc.add_paragraph()  

    # Write Diagram Questions (At the very end)
    if exam.diagram_questions:
        doc.add_paragraph() 
        diag_header = doc.add_paragraph()
        diag_header.add_run("Diagrams & Visuals").bold = True
        
        for i, dq in enumerate(exam.diagram_questions, 1):
            p = doc.add_paragraph()
            q_text = f"{i}. {dq.question} ({dq.marks} Marks)"
            if show_clo and dq.target_clo:
                q_text += f" [{dq.target_clo}]"
            p.add_run(q_text).bold = True
            insert_image_if_exists(dq.image_url)
            doc.add_paragraph()

    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0) 
    return output_stream


# --- 3. FASTAPI ENDPOINT ---
@app.post("/api/v1/generate-document")
async def generate_document(request_data: DocumentRequest):
    try:
        final_doc_stream = process_exam(request_data)
        headers = {
            "Content-Disposition": "attachment; filename=generated_exam.docx"
        }
        return StreamingResponse(
            final_doc_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))